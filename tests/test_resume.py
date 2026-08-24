from pathlib import Path

import docx
import pytest

from lvtest.errors import LvtestError
from lvtest.resume import extract_text


def _write_pdf(path: Path, text: str | None) -> None:
    """xref 오프셋을 정확히 계산한 최소 PDF. text=None이면 내용 없는 페이지(스캔본 흉내)."""
    content = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET" if text is not None else ""
    objs = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R "
        "/Resources << /Font << /F1 5 0 R >> >> >>",
        f"<< /Length {len(content)} >>\nstream\n{content}\nendstream",
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, body in enumerate(objs, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n{body}\nendobj\n".encode("latin-1")
    xref_pos = len(out)
    out += f"xref\n0 {len(objs) + 1}\n0000000000 65535 f \n".encode()
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += f"trailer\n<< /Size {len(objs) + 1} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode()
    path.write_bytes(bytes(out))


def test_markdown(tmp_path):
    p = tmp_path / "resume.md"
    p.write_text("# 홍길동\n\n- Spring Boot 주문 서비스 개발\n", encoding="utf-8")
    assert "Spring Boot 주문 서비스 개발" in extract_text(p)


def test_pdf_with_text(tmp_path):
    p = tmp_path / "resume.pdf"
    _write_pdf(p, "Backend Engineer Resume")
    assert "Backend Engineer Resume" in extract_text(p)


def test_pdf_without_text_is_unreadable(tmp_path):
    p = tmp_path / "scan.pdf"
    _write_pdf(p, None)
    with pytest.raises(LvtestError) as ei:
        extract_text(p)
    assert ei.value.code == "resume_unreadable"
    assert "md" in ei.value.message


def test_docx_paragraphs_and_tables(tmp_path):
    p = tmp_path / "resume.docx"
    d = docx.Document()
    d.add_paragraph("김개발 — 백엔드 3년")
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "기술"
    table.rows[0].cells[1].text = "Kotlin, PostgreSQL"
    d.save(p)
    text = extract_text(p)
    assert "김개발 — 백엔드 3년" in text
    assert "Kotlin, PostgreSQL" in text


def test_unsupported_extension(tmp_path):
    p = tmp_path / "resume.hwp"
    p.write_bytes(b"x")
    with pytest.raises(LvtestError) as ei:
        extract_text(p)
    assert ei.value.code == "resume_unreadable"


def test_missing_file(tmp_path):
    with pytest.raises(LvtestError) as ei:
        extract_text(tmp_path / "nope.md")
    assert ei.value.code == "resume_unreadable"
