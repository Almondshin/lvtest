from __future__ import annotations

from pathlib import Path

from lvtest.errors import LvtestError

TEXT_SUFFIXES = {".md", ".markdown", ".txt"}
HINT = "텍스트가 있는 pdf/docx/md 파일이어야 합니다. 스캔본이면 md로 변환해서 다시 시도하세요."


def _unreadable(path: Path, why: str) -> LvtestError:
    return LvtestError("resume_unreadable", f"{why} ({path.name}). {HINT}", path=str(path))


def _from_pdf(path: Path) -> str:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            raise _unreadable(path, "암호화된 PDF입니다")
        pages = [page.extract_text() or "" for page in reader.pages]
    except PdfReadError as e:
        raise _unreadable(path, f"PDF를 읽을 수 없습니다: {e}") from e
    return "\n\n".join(p.strip() for p in pages if p.strip())


def _from_docx(path: Path) -> str:
    import docx

    try:
        document = docx.Document(str(path))
    except Exception as e:  # Parser boundary: docx/lxml/zipfile can raise various exceptions
        raise _unreadable(path, f"DOCX를 읽을 수 없습니다: {e}") from e
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(path: Path) -> str:
    path = Path(path)
    if not path.is_file():
        raise _unreadable(path, "파일이 없습니다")
    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        text = path.read_text(encoding="utf-8", errors="replace")
    elif suffix == ".pdf":
        text = _from_pdf(path)
    elif suffix == ".docx":
        text = _from_docx(path)
    else:
        raise _unreadable(path, f"지원하지 않는 확장자 '{suffix}'")
    if not text.strip():
        raise _unreadable(path, "추출된 텍스트가 없습니다 (스캔본?)")
    return text.strip()
